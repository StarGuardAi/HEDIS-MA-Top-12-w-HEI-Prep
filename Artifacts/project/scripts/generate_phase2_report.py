"""
Generate Phase 2 Completion Report as .docx
Run from repo root or Artifacts/project/
"""
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
except ImportError:
    print("Installing python-docx...")
    import subprocess
    subprocess.check_call(["pip", "install", "python-docx"])
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE

def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    return p

def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if bold:
        run.bold = True
    return p

def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hrow = t.rows[0]
    for i, h in enumerate(headers):
        hrow.cells[i].text = h
        for p in hrow.cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        r = t.rows[ri + 1]
        for ci, cell in enumerate(row):
            if ci < len(r.cells):
                r.cells[ci].text = str(cell)
    if col_widths:
        for i, w in enumerate(col_widths):
            if i < len(t.columns):
                for cell in t.columns[i].cells:
                    cell.width = Inches(w)
    return t

def main():
    out_dir = Path(__file__).resolve().parent.parent
    out_path = out_dir / "PHASE2_COMPLETION_REPORT.docx"

    doc = Document()
    doc.add_heading("Phase 2 Completion Report", 0)
    doc.add_paragraph("Hardening Sprint Gate — AuditShield Portfolio")
    doc.add_paragraph()

    # Executive Summary
    add_heading(doc, "Executive Summary", 1)
    add_para(doc, "Sprint scope: Active suppression, HITL Admin View, Intervention Optimizer, and hardening artifacts (pyproject.toml, tests, docs) across AuditShield, StarGuard Desktop, and StarGuard Mobile.")
    add_para(doc, "Current gate status: OPEN", bold=True)
    add_para(doc, "Gate closes when all six Phase 2 criteria are met for each app.")
    doc.add_paragraph()

    # Portfolio Overview
    add_heading(doc, "Portfolio Overview", 1)
    add_table(doc,
        ["App", "HuggingFace URL", "Local Port"],
        [
            ["AuditShield", "https://huggingface.co/spaces/rreichert/auditshield-live", "8501"],
            ["StarGuard Desktop", "https://rreichert-starguard-desktop.hf.space", "8000"],
            ["StarGuard Mobile", "https://rreichert-starguardai.hf.space", "8000"],
        ],
        col_widths=[1.2, 2.8, 0.8]
    )
    doc.add_paragraph()

    # Task Definitions
    add_heading(doc, "Task Definitions", 1)
    tasks = [
        ("T1", "pyproject.toml", "Authoritative build config, pinned deps, pip install -e . works"),
        ("T2", "Type hints & separation", "Public functions typed; business logic separated from Shiny reactive"),
        ("T3", "Unit tests", "Min 3 tests per new module; pytest runs clean"),
        ("T4", "Documentation", "ARCHITECTURE.md, README current"),
        ("T5", "Pre-phase gate", "No embedded refactor; tests before sign-off"),
        ("T6", "starguard-core [Phase 3]", "Shared package extraction — deferred until Phase 2 gate closes"),
    ]
    add_table(doc, ["Task", "Deliverable", "Definition"], tasks, col_widths=[0.6, 1.5, 3.5])
    doc.add_paragraph()

    # Per-App Status Tables
    add_heading(doc, "Per-App Status Tables", 1)
    status_headers = ["Artifact", "AuditShield", "Desktop", "Mobile"]
    status_rows = [
        ("pyproject.toml", "⏳ PENDING", "⏳ PENDING", "⏳ PENDING"),
        ("Type hints & separation", "⏳ PENDING", "⏳ PENDING", "⏳ PENDING"),
        ("Unit tests (min 3/module)", "⏳ PENDING", "⏳ PENDING", "⏳ PENDING"),
        ("ARCHITECTURE.md / README", "⏳ PENDING", "⏳ PENDING", "⏳ PENDING"),
        ("Pre-phase gate", "⏳ PENDING", "⏳ PENDING", "⏳ PENDING"),
    ]
    add_table(doc, status_headers, status_rows, col_widths=[1.8, 1.2, 1.2, 1.2])
    doc.add_paragraph()

    # Phase 2 Gate Criteria
    add_heading(doc, "Phase 2 Gate Criteria", 1)
    add_para(doc, "All six conditions must be true to close the gate:")
    for i, c in enumerate([
        "pyproject.toml exists and is authoritative in each app",
        "Dependencies pinned with versions",
        "pip install -e . (or equivalent) works from each repo root",
        "All public functions have type hints; business logic separated from Shiny",
        "Minimum 3 unit tests per new module; pytest runs clean",
        "ARCHITECTURE.md and README current; no embedded refactoring",
    ], 1):
        add_para(doc, f"  {i}. {c}")
    doc.add_paragraph()

    # Phase 3 Deferral
    add_heading(doc, "Phase 3 Deferral", 1)
    add_para(doc, "starguard-core: Shared package extraction (hedis_gap_trail, hitl_admin_view, suppression_banner, etc.) is deferred to Phase 3. Rationale: Phase 2 hardening must complete first to establish consistent structure, tests, and docs across all three apps. starguard-core depends on Phase 2 gate close.")
    doc.add_paragraph()

    # Configuration Reference
    add_heading(doc, "Configuration Reference", 1)
    add_table(doc, ["Item", "Path / Value"],
        [
            ("cursorrules", "Artifacts/project/auditshield/.cursorrules (and per-app)"),
            ("Docker", "Per-app Dockerfile in each repo"),
            ("Ports", "AuditShield: 8501; Desktop/Mobile: 8000"),
            ("Org name", "rreichert (HuggingFace)"),
        ],
        col_widths=[1.5, 3.5]
    )
    doc.add_paragraph()

    # Commit Instructions
    add_heading(doc, "Commit Instructions", 1)
    add_para(doc, "Report path: Artifacts/project/PHASE2_COMPLETION_REPORT.docx")
    add_para(doc, "Commit message: chore: add Phase 2 Completion Report (hardening sprint gate)")
    add_para(doc, "Repos: Main HEDIS-MA-Top-12-w-HEI-Prep (report lives here). Desktop and Mobile are submodules with separate repos.")
    add_para(doc, "When the five hardening artifacts are generated per app, update the status tables to ✓ DONE and re-commit as the gate-close record.")
    doc.add_paragraph()
    add_para(doc, "Ready to start generating the actual artifacts (pyproject.toml first) whenever you are.")

    doc.save(out_path)
    print(f"Created: {out_path}")

if __name__ == "__main__":
    main()
