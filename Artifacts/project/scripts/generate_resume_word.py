#!/usr/bin/env python3
"""
One-Page Word Resume Generator for HEDIS GSD Project

Automatically generates a professional one-page resume in Word format (.docx)
highlighting your completed milestones and achievements.

Requirements:
    pip install python-docx

Usage:
    python scripts/generate_resume_word.py
    python scripts/generate_resume_word.py --milestone 1
    python scripts/generate_resume_word.py --style modern
"""

import os
import sys
import json
import argparse
from pathlib import Path
from datetime import datetime
import logging

try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    Document = None
    print("ERROR: python-docx not installed")
    print("Install with: pip install python-docx")
    sys.exit(1)

logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)


class ResumeGenerator:
    """Generates professional one-page Word resume"""
    
    def __init__(self):
        self.project_root = Path(__file__).parent.parent
        self.milestones = self._load_milestones()
        self.doc = Document()
        
        # Configure page margins for one-page layout (tighter margins)
        sections = self.doc.sections
        for section in sections:
            section.top_margin = Inches(0.4)  # Reduced from 0.5
            section.bottom_margin = Inches(0.4)  # Reduced from 0.5
            section.left_margin = Inches(0.6)  # Reduced from 0.7
            section.right_margin = Inches(0.6)  # Reduced from 0.7
    
    def _load_milestones(self) -> dict:
        """Load milestone data"""
        milestone_file = self.project_root / 'milestones.json'
        
        if milestone_file.exists():
            with open(milestone_file, 'r', encoding='utf-8') as f:
                return json.load(f)
        return {}
    
    def add_header(self, name: str = "Robert Reichert",
                   title: str = "Healthcare Data Scientist & AI Engineer",
                   contact: dict = None):
        """Add resume header"""
        if contact is None:
            contact = {
                'location': 'Pittsburgh, PA',
                'phone': '480-767-1337',
                'email': 'reichert.starguardai@gmail.com',
                'linkedin': 'linkedin.com/in/rreichert-hedis-data-science-ai',
                'github': 'github.com/StarGuardAi',
                'portfolio': 'Canva Portfolio'
            }
        
        # Name (compact for single-page)
        name_para = self.doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_para.space_after = Pt(2)  # Reduced spacing
        name_run = name_para.add_run(name.upper())
        name_run.bold = True
        name_run.font.size = Pt(16)  # Reduced from 18
        name_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
        
        # Title
        title_para = self.doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.space_after = Pt(2)  # Reduced spacing
        title_run = title_para.add_run(title)
        title_run.font.size = Pt(10)  # Reduced from 11
        title_run.font.color.rgb = RGBColor(51, 51, 51)  # Dark gray
        
        # Contact info (compact for single-page layout)
        contact_para = self.doc.add_paragraph()
        contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        contact_para.space_after = Pt(6)  # Reduced spacing
        contact_text = f"{contact['email']} • {contact['linkedin']} • {contact['github']}"
        
        contact_run = contact_para.add_run(contact_text)
        contact_run.font.size = Pt(9)
        
        # Add line separator
        self.doc.add_paragraph('_' * 100).runs[0].font.size = Pt(8)
    
    def add_section(self, title: str):
        """Add section header (compact for single-page)"""
        section_para = self.doc.add_paragraph()
        section_para.space_before = Pt(4)  # Reduced spacing
        section_para.space_after = Pt(2)  # Reduced spacing
        section_run = section_para.add_run(title.upper())
        section_run.bold = True
        section_run.font.size = Pt(11)  # Reduced from 12
        section_run.font.color.rgb = RGBColor(0, 51, 102)  # Dark blue
    
    def add_summary(self):
        """Add professional summary (compact)"""
        summary_text = """Medicare Advantage Diabetes Analytics Specialist with expertise in HEDIS gap closure prediction. Focus on diabetes portfolio (HBD, KED, EED, PDC-DR) - four measures worth $620K-$1M in Star revenue protection. Using AI-assisted development (Cursor AI, Claude Sonnet, ChatGPT-4), I deliver predictive models in 60-90 days vs. 6-12 months traditional consulting. Recent project: 91% AUC-ROC predicting poor glycemic control with 15-25% improvement potential. Available for part-time/fractional projects."""
        
        para = self.doc.add_paragraph(summary_text)
        para.paragraph_format.space_after = Pt(4)  # Reduced from 6
        para.paragraph_format.space_before = Pt(2)
        para.runs[0].font.size = Pt(9.5)  # Reduced slightly
    
    def add_featured_project(self, milestone_id: int = None):
        """Add featured project (HEDIS GSD Prediction Engine)"""
        # Get completed milestones
        completed = [m for m in self.milestones.get('milestones', []) if m['status'] == 'completed']
        
        if not completed:
            logger.warning("No completed milestones found")
            return
        
        # Project header
        project_para = self.doc.add_paragraph()
        project_run = project_para.add_run("HEALTHCARE ML ENGINEER | ")
        project_run.bold = True
        project_run.font.size = Pt(10.5)
        
        title_run = project_para.add_run("HEDIS GSD Prediction Engine | ")
        title_run.font.size = Pt(10.5)
        
        date_run = project_para.add_run("2025")
        date_run.font.size = Pt(10.5)
        date_run.italic = True
        
        # AI Tools emphasis
        ai_para = self.doc.add_paragraph()
        ai_run = ai_para.add_run("AI-Assisted Development with Cursor AI, Claude Sonnet 3.5, ChatGPT-4")
        ai_run.font.size = Pt(10)
        ai_run.italic = True
        ai_run.font.color.rgb = RGBColor(51, 51, 51)
        
        # Key achievements
        achievements = [
            "Built diabetes portfolio prediction models targeting four HEDIS measures (HBD, KED, EED, PDC-DR) worth $620K-$1.08M combined value for 100K-member MA plan, achieving 91% AUC-ROC for glycemic control prediction and demonstrating 15-25% improvement potential in gap closure rates",
            
            "Pioneered AI-assisted development using Cursor AI + Claude Sonnet + ChatGPT-4, delivering production-ready predictive models in 60-90 days vs. 6-12 months traditional consulting, achieving 60% faster development at 10x lower cost while maintaining 100% HIPAA compliance",
            
            "Engineered 25+ HEDIS-compliant features focusing on diabetes portfolio measures, processing 150,000+ Medicare claims for 24,935 diabetic members with automated de-identification and PHI-safe logging, targeting triple-weighted measures (HBD, KED) for maximum Star Ratings impact",
            
            "Implemented SHAP interpretability with portfolio-level risk stratification, identifying 6,200+ high-risk members across diabetes measures and enabling targeted interventions for HbA1c control, kidney evaluation, eye exams, and medication adherence"
        ]
        
        for achievement in achievements:
            bullet_para = self.doc.add_paragraph(achievement, style='List Bullet')
            bullet_para.paragraph_format.space_after = Pt(2)  # Reduced from 3
            bullet_para.paragraph_format.space_before = Pt(0)
            bullet_para.paragraph_format.left_indent = Inches(0.2)  # Reduced indent
            bullet_para.paragraph_format.line_spacing = 1.0  # Single spacing
            for run in bullet_para.runs:
                run.font.size = Pt(9.5)  # Reduced from 10
        
        # Technologies (compact)
        tech_para = self.doc.add_paragraph()
        tech_para.space_after = Pt(4)  # Reduced from 6
        tech_label = tech_para.add_run("Technologies: ")
        tech_label.bold = True
        tech_label.font.size = Pt(9)  # Reduced
        
        tech_text = tech_para.add_run("Cursor AI, Claude Sonnet, ChatGPT-4, Python, scikit-learn, XGBoost, SHAP, pandas, HEDIS (HBD/KED/EED/PDC-DR), Medicare Advantage, Star Ratings, HIPAA")
        tech_text.font.size = Pt(9)  # Reduced
        tech_para.paragraph_format.line_spacing = 1.0  # Single spacing
    
    def add_milestones_summary(self):
        """Add milestones completion summary"""
        completed = [m for m in self.milestones.get('milestones', []) if m['status'] == 'completed']
        in_progress = [m for m in self.milestones.get('milestones', []) if m['status'] == 'in_progress']
        
        if not completed:
            return
        
        # Milestones header
        milestone_para = self.doc.add_paragraph()
        milestone_run = milestone_para.add_run("Development Milestones Completed:")
        milestone_run.bold = True
        milestone_run.font.size = Pt(10)
        
        # List completed milestones
        for milestone in completed:
            bullet = f"✅ Milestone {milestone['id']}: {milestone['title']} ({milestone.get('completion_date', 'TBD')})"
            bullet_para = self.doc.add_paragraph(bullet, style='List Bullet')
            bullet_para.paragraph_format.space_after = Pt(2)
            bullet_para.paragraph_format.left_indent = Inches(0.25)
            for run in bullet_para.runs:
                run.font.size = Pt(9)
        
        # In progress
        if in_progress:
            for milestone in in_progress:
                bullet = f"🔄 Milestone {milestone['id']}: {milestone['title']} - IN PROGRESS"
                bullet_para = self.doc.add_paragraph(bullet, style='List Bullet')
                bullet_para.paragraph_format.space_after = Pt(2)
                bullet_para.paragraph_format.left_indent = Inches(0.25)
                for run in bullet_para.runs:
                    run.font.size = Pt(9)
    
    def add_skills(self):
        """Add skills section"""
        skills = {
            'AI-Assisted Development': 'Cursor AI IDE, Claude Sonnet, ChatGPT, AI-powered code generation and review',
            'Healthcare Analytics': 'HEDIS Quality Measures, HIPAA Compliance, CMS Medicare Data, Clinical Validation',
            'Machine Learning': 'scikit-learn, Predictive Modeling, SHAP Interpretability, Bias Analysis, Temporal Validation',
            'Data Engineering': 'Python, pandas, NumPy, ETL Pipelines, Feature Engineering, Data Quality',
            'Software Engineering': 'pytest, Git, CI/CD, Production MLOps, Comprehensive Testing, Documentation'
        }
        
        for category, skills_list in skills.items():
            skill_para = self.doc.add_paragraph()
            skill_para.paragraph_format.space_after = Pt(2)
            
            category_run = skill_para.add_run(f"{category}: ")
            category_run.bold = True
            category_run.font.size = Pt(9.5)
            
            skills_run = skill_para.add_run(skills_list)
            skills_run.font.size = Pt(9.5)
    
    def add_key_metrics(self):
        """Add key metrics/achievements box"""
        # Create a bordered box effect with table
        table = self.doc.add_table(rows=1, cols=4)
        table.style = 'Light Grid Accent 1'
        
        metrics = [
            ('91% AUC-ROC', 'Model Performance'),
            ('100% HIPAA', 'Compliance'),
            ('4,800+ Lines', 'Production Code'),
            ('60% Faster', 'with AI Tools')
        ]
        
        for i, (value, label) in enumerate(metrics):
            cell = table.rows[0].cells[i]
            cell.text = f"{value}\n{label}"
            
            # Center align
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.size = Pt(9)
                    if value in run.text:
                        run.bold = True
    
    def generate(self, output_filename: str = None, style: str = 'modern') -> Path:
        """
        Generate complete resume
        
        Args:
            output_filename: Output filename (default: auto-generated)
            style: Resume style ('modern' or 'classic')
        
        Returns:
            Path to generated resume
        """
        logger.info(f"Generating {style} resume...")
        
        # Header
        self.add_header()
        
        # Professional Summary
        self.add_section("Professional Summary")
        self.add_summary()
        
        # Featured Project
        self.add_section("Featured Project")
        self.add_featured_project()
        
        # Milestones (if completed)
        if any(m['status'] == 'completed' for m in self.milestones.get('milestones', [])):
            self.add_milestones_summary()
        
        # Skills
        self.add_section("Technical Skills")
        self.add_skills()
        
        # Key Metrics
        self.doc.add_paragraph()  # Spacing
        self.add_key_metrics()
        
        # Save document
        if not output_filename:
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_filename = f"Resume_HEDIS_GSD_{timestamp}.docx"
        
        output_path = self.project_root / 'reports' / output_filename
        output_path.parent.mkdir(exist_ok=True)
        
        self.doc.save(str(output_path))
        
        logger.info(f"✅ Resume generated: {output_path}")
        return output_path
    
    def _update_publishing_status(self, milestone_id: int):
        """Update resume publishing status"""
        milestone_file = self.project_root / 'milestones.json'
        
        with open(milestone_file, 'r', encoding='utf-8') as f:
            data = json.load(f)
        
        for milestone in data.get('milestones', []):
            if milestone['id'] == milestone_id:
                milestone['publishing_status']['resume'] = 'published'
                break
        
        # Update last_updated timestamp
        data['last_updated'] = datetime.now().isoformat()
        
        with open(milestone_file, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2)
        
        logger.info(f"Updated resume status to 'published' for Milestone {milestone_id}")


def main():
    """Main function"""
    parser = argparse.ArgumentParser(
        description='Generate one-page Word resume for HEDIS GSD Project'
    )
    parser.add_argument(
        '--milestone',
        type=int,
        help='Mark specific milestone as published in resume'
    )
    parser.add_argument(
        '--output',
        help='Output filename (default: auto-generated with timestamp)'
    )
    parser.add_argument(
        '--style',
        choices=['modern', 'classic'],
        default='modern',
        help='Resume style'
    )
    
    args = parser.parse_args()
    
    # Generate resume
    generator = ResumeGenerator()
    output_path = generator.generate(
        output_filename=args.output,
        style=args.style
    )
    
    # Update milestone status if specified
    if args.milestone:
        generator._update_publishing_status(args.milestone)
    
    print(f"\n✅ Resume generated successfully!")
    print(f"📄 Location: {output_path}")
    print(f"\n💡 Next steps:")
    print(f"   1. Open the resume in Microsoft Word")
    print(f"   2. Review and customize as needed")
    print(f"   3. Save as PDF for applications")
    print(f"   4. Update with your actual contact information")
    
    # Try to open the file (Windows only)
    if os.name == 'nt':
        try:
            os.startfile(output_path)
            print(f"\n📂 Opening resume in Word...")
        except:
            pass
    
    return 0


if __name__ == "__main__":
    sys.exit(main())


