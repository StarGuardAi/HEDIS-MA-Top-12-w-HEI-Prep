"""Update email and portfolio link in resume DOCX file"""
import docx
from docx import Document
import os
from datetime import datetime

def update_resume():
    """Update email and portfolio link in the latest resume"""
    
    # Input file
    input_file = "reports/Robert_Reichert_Resume_20251026_171243.docx"
    
    # Output file with timestamp
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    output_file = f"reports/Robert_Reichert_Resume_{timestamp}.docx"
    
    print(f"Loading resume: {input_file}")
    doc = Document(input_file)
    
    # Update all text in paragraphs and tables
    changes_made = 0
    
    # Update in paragraphs
    for paragraph in doc.paragraphs:
        if "reichert.starguardai@gmail.com" in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace("reichert.starguardai@gmail.com", "robert.reichert.starguardai@gmail.com")
                changes_made += 1
        
        if "bobareichert.github.io/HEDIS-MA-Top-12-w-HEI-Prep" in paragraph.text:
            for run in paragraph.runs:
                run.text = run.text.replace(
                    "bobareichert.github.io/HEDIS-MA-Top-12-w-HEI-Prep",
                    "bobareichert.github.io/HEDIS-MA-Top-12-w-HEI-Prep/"
                )
                changes_made += 1
    
    # Update in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if "reichert.starguardai@gmail.com" in paragraph.text:
                        for run in paragraph.runs:
                            run.text = run.text.replace("reichert.starguardai@gmail.com", "robert.reichert.starguardai@gmail.com")
                            changes_made += 1
                    
                    if "bobareichert.github.io/HEDIS-MA-Top-12-w-HEI-Prep" in paragraph.text:
                        for run in paragraph.runs:
                            run.text = run.text.replace(
                                "bobareichert.github.io/HEDIS-MA-Top-12-w-HEI-Prep",
                                "bobareichert.github.io/HEDIS-MA-Top-12-w-HEI-Prep/"
                            )
                            changes_made += 1
    
    # Save updated resume
    print(f"Saving updated resume: {output_file}")
    doc.save(output_file)
    
    print(f"\n✅ Resume updated successfully!")
    print(f"Changes made: {changes_made}")
    print(f"Output file: {output_file}")
    
    # Also create a copy as the "latest" version
    latest_file = "reports/Robert_Reichert_Resume_LATEST.docx"
    doc.save(latest_file)
    print(f"Saved as: {latest_file}")
    
    return output_file, latest_file

if __name__ == "__main__":
    update_resume()

